"""Tests for the Bid Decision Report (reasoning trace)."""
from docx import Document

from tools.bid_report import append_report_to_docx, build_report

EVIDENCE_MAP = {
    "REQ-001": [
        {"doc_id": "DOC-001", "title": "HealthGov Case Study", "excerpt": "...",
         "source_path": "corpus/DOC-001.md", "score": 15.0},
    ],
    "REQ-002": [],
}

VERIFIED_DRAFT = {
    "rfp_title": "Test RFP",
    "coverage_score": 60,
    "gap_count": 1,
    "verification": {"citations_total": 1, "citations_verified": 1, "citations_stripped": 0},
    "requirements": [
        {
            "id": "REQ-001", "text": "Migration experience required.",
            "priority": "high", "category": "technical", "confidence": 0.9,
            "score": "COVERED", "response_text": "Grounded [DOC-001].",
            "evidence_citations": "[DOC-001] HealthGov Case Study", "gap_note": None,
            "verification": {"cited": ["DOC-001"], "verified": ["DOC-001"], "stripped": []},
        },
        {
            "id": "REQ-002", "text": "Quantum-safe cryptography plan required.",
            "priority": "medium", "category": "technical", "confidence": 0.95,
            "score": "GAP", "response_text": None, "evidence_citations": None,
            "gap_note": "Provide a quantum-safe cryptography transition plan.",
        },
    ],
}


def test_report_structure_and_decision_chain():
    report = build_report(VERIFIED_DRAFT, EVIDENCE_MAP)
    assert report["coverage_score"] == 60
    assert report["recommendation"] == "BID WITH CONDITIONS"
    assert report["counts"] == {"total": 2, "covered": 1, "partial": 0, "gap": 1}
    assert report["citations"]["verified"] == 1

    covered = report["requirements"][0]
    assert covered["evidence_considered"][0]["doc_id"] == "DOC-001"
    assert covered["decision"]["score"] == "COVERED"
    assert covered["decision"]["scorer_confidence"] == 0.9
    assert covered["decision"]["citations_verified"] == ["DOC-001"]
    assert covered["action_required"] is None

    gap = report["requirements"][1]
    assert gap["evidence_considered"] == []
    assert gap["decision"]["score"] == "GAP"
    assert gap["action_required"].startswith("Provide a quantum-safe")


def test_recommendation_bands():
    high = dict(VERIFIED_DRAFT, coverage_score=85)
    low = dict(VERIFIED_DRAFT, coverage_score=30)
    assert build_report(high, EVIDENCE_MAP)["recommendation"] == "BID"
    assert build_report(low, EVIDENCE_MAP)["recommendation"] == "REVIEW BID DECISION"


def test_report_appended_to_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Proposal body")
    docx_path = str(tmp_path / "proposal.docx")
    doc.save(docx_path)

    report = build_report(VERIFIED_DRAFT, EVIDENCE_MAP)
    append_report_to_docx(docx_path, report)

    text = "\n".join(p.text for p in Document(docx_path).paragraphs)
    assert "Appendix — Bid Decision Report" in text
    assert "BID WITH CONDITIONS" in text
    assert "REQ-001" in text and "REQ-002" in text
    assert "DOC-001 (HealthGov Case Study)" in text
    assert "ACTION REQUIRED: Provide a quantum-safe" in text
    assert "no relevant documents in the knowledge base" in text
