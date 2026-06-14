"""DOCX proposal builder — fills the proposal template with pipeline output data."""
import copy
import os
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_proposal(data: dict, template_path: str, output_path: str) -> str:
    """Fill the proposal template with data and save to output_path.

    Args:
        data: dict with keys:
            company_name, rfp_title, submission_date, executive_summary,
            coverage_score, gap_summary,
            requirements: [{"id", "text", "score", "response_text", "evidence_citations", "gap_note"}]
        template_path: path to proposal_template.docx
        output_path: where to write the populated .docx

    Returns:
        Absolute path of the written file.
    """
    doc = Document(template_path)

    simple_replacements = {
        "{{company_name}}": str(data.get("company_name", "")),
        "{{rfp_title}}": str(data.get("rfp_title", "")),
        "{{submission_date}}": str(data.get("submission_date", "")),
        "{{executive_summary}}": str(data.get("executive_summary", "")),
        "{{coverage_score}}": str(data.get("coverage_score", "")),
        "{{gap_summary}}": _build_gap_summary(data),
    }

    _replace_in_paragraphs(doc, simple_replacements)
    _fill_requirements_table(doc, data.get("requirements", []))

    os.makedirs(Path(output_path).parent, exist_ok=True)
    doc.save(output_path)
    return str(Path(output_path).resolve())


def _replace_in_paragraphs(doc: Document, replacements: dict) -> None:
    """Replace {{placeholder}} tokens in all paragraphs (including headers/footers)."""
    all_paragraphs = list(doc.paragraphs)
    for section in doc.sections:
        all_paragraphs += list(section.header.paragraphs)
        all_paragraphs += list(section.footer.paragraphs)

    for para in all_paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                _replace_in_para(para, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in para.text:
                            _replace_in_para(para, key, value)


def _replace_in_para(para, placeholder: str, value: str) -> None:
    """Replace a placeholder within a paragraph's runs, preserving formatting."""
    full_text = para.text
    if placeholder not in full_text:
        return

    new_text = full_text.replace(placeholder, value)

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)


def _fill_requirements_table(doc: Document, requirements: list) -> None:
    """Find the requirements table and replace placeholder rows with real data."""
    for table in doc.tables:
        if _is_requirements_table(table):
            _rebuild_requirements_table(table, requirements)
            return


def _is_requirements_table(table) -> bool:
    """Detect the requirements table by its header row content."""
    if not table.rows:
        return False
    header_text = " ".join(cell.text for cell in table.rows[0].cells)
    return "Req ID" in header_text or "Our Response" in header_text


def _rebuild_requirements_table(table, requirements: list) -> None:
    """Clear placeholder rows and insert one row per requirement."""
    header_row = table.rows[0]

    # Remove all rows except the header
    for row in list(table.rows)[1:]:
        tr = row._tr
        tr.getparent().remove(tr)

    for req in requirements:
        is_gap = req.get("score") == "GAP"
        _add_requirement_row(table, req, is_gap)


def _add_requirement_row(table, req: dict, is_gap: bool) -> None:
    row = table.add_row()
    cells = row.cells

    req_id = req.get("id", "")
    response_text = (
        f"[ACTION REQUIRED: {req.get('gap_note', 'No gap note provided')}]"
        if is_gap
        else req.get("response_text", "")
    )
    evidence_text = (
        "No evidence on file — human action required."
        if is_gap
        else req.get("evidence_citations", "")
    )

    cell_values = [req_id, response_text, evidence_text]

    for i, (cell, value) in enumerate(zip(cells, cell_values)):
        para = cell.paragraphs[0]
        run = para.add_run(value)
        run.font.size = Pt(10)
        run.font.name = "Arial"

        if is_gap:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            _set_cell_shading(cell, "FFE0E0")
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _set_cell_shading(cell, fill_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def _build_gap_summary(data: dict) -> str:
    requirements = data.get("requirements", [])
    gaps = [r for r in requirements if r.get("score") == "GAP"]
    partials = [r for r in requirements if r.get("score") == "PARTIAL"]

    if not gaps and not partials:
        return "No gaps identified. All requirements are fully covered by internal evidence."

    lines = []
    if gaps:
        lines.append(f"GAPS ({len(gaps)} items requiring human input before submission):")
        for r in gaps:
            lines.append(f"  [ACTION REQUIRED] {r['id']}: {r.get('gap_note', 'See requirement details.')}")
    if partials:
        lines.append(f"\nPARTIAL COVERAGE ({len(partials)} items with weak evidence):")
        for r in partials:
            lines.append(f"  {r['id']}: {r.get('gap_note', 'Additional evidence recommended.')}")

    return "\n".join(lines)
