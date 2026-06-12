import logging
import os
import re
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)


def _stub_mode() -> bool:
    return os.getenv("STUB_MODE", "false").lower() == "true"


def parse_rfp(file_path: str) -> dict:
    """Parse a PDF or DOCX RFP and return structured text.

    Returns:
        {
            "full_text": str,
            "pages": int,
            "sections": [{"heading": str, "content": str}]
        }
    """
    # STUB: return hardcoded sample when STUB_MODE=true
    # TODO: remove stub once live Azure credentials are configured
    if _stub_mode():
        return _stub_result()

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"RFP file not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _parse_docx(path)
    elif suffix == ".pdf":
        return _parse_pdf_with_document_intelligence(path)
    else:
        raise ValueError(f"Unsupported file type '{suffix}'. Expected .pdf or .docx.")


def _parse_pdf_with_document_intelligence(path: Path) -> dict:
    endpoint = os.getenv("DOC_INTELLIGENCE_ENDPOINT")
    key = os.getenv("DOC_INTELLIGENCE_KEY")

    if not endpoint or not key:
        logger.warning(
            "DOC_INTELLIGENCE_ENDPOINT / DOC_INTELLIGENCE_KEY not set — falling back "
            "to local pypdf parsing. Azure Document Intelligence is the documented "
            "path and gives higher-fidelity layout/section extraction."
        )
        return _parse_pdf_with_pypdf(path)

    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential
        from azure.core.exceptions import HttpResponseError
    except ImportError as e:
        raise ImportError(
            "azure-ai-documentintelligence is required. Run: pip install azure-ai-documentintelligence"
        ) from e

    client = DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))

    try:
        with open(path, "rb") as f:
            poller = client.begin_analyze_document("prebuilt-layout", body=f, content_type="application/octet-stream")
        result = poller.result()
    except Exception as e:
        raise RuntimeError(f"Azure Document Intelligence error: {e}") from e

    full_text_parts = []
    sections = []
    current_heading = "Introduction"
    current_content_parts = []

    for page in result.pages or []:
        for line in page.lines or []:
            full_text_parts.append(line.content)

    for paragraph in result.paragraphs or []:
        role = getattr(paragraph, "role", None)
        content = paragraph.content or ""
        if role in ("title", "sectionHeading"):
            if current_content_parts:
                sections.append({
                    "heading": current_heading,
                    "content": " ".join(current_content_parts).strip(),
                })
                current_content_parts = []
            current_heading = content
        else:
            current_content_parts.append(content)

    if current_content_parts:
        sections.append({
            "heading": current_heading,
            "content": " ".join(current_content_parts).strip(),
        })

    return {
        "full_text": "\n".join(full_text_parts),
        "pages": len(result.pages or []),
        "sections": sections,
    }


_PYPDF_HEADING_RE = re.compile(r"^(?:\d+\.\s+\S.*|Appendix\s+[A-Z]\b.*)$")


def _parse_pdf_with_pypdf(path: Path) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("pypdf is required for local PDF parsing. Run: pip install pypdf") from e

    reader = PdfReader(str(path))
    full_text_parts = []
    sections = []
    current_heading = "Introduction"
    current_content_parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            full_text_parts.append(line)

            if _PYPDF_HEADING_RE.match(line):
                if current_content_parts:
                    sections.append({
                        "heading": current_heading,
                        "content": " ".join(current_content_parts).strip(),
                    })
                    current_content_parts = []
                current_heading = line
            else:
                current_content_parts.append(line)

    if current_content_parts:
        sections.append({
            "heading": current_heading,
            "content": " ".join(current_content_parts).strip(),
        })

    return {
        "full_text": "\n".join(full_text_parts),
        "pages": len(reader.pages),
        "sections": sections,
    }


def _parse_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required. Run: pip install python-docx") from e

    doc = Document(str(path))
    full_text_parts = []
    sections = []
    current_heading = "Introduction"
    current_content_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text_parts.append(text)

        if para.style.name.startswith("Heading"):
            if current_content_parts:
                sections.append({
                    "heading": current_heading,
                    "content": " ".join(current_content_parts).strip(),
                })
                current_content_parts = []
            current_heading = text
        else:
            current_content_parts.append(text)

    if current_content_parts:
        sections.append({
            "heading": current_heading,
            "content": " ".join(current_content_parts).strip(),
        })

    return {
        "full_text": "\n".join(full_text_parts),
        "pages": len(doc.sections),
        "sections": sections,
    }


def _stub_result() -> dict:
    return {
        "full_text": (
            "REQUEST FOR PROPOSAL\n"
            "Cloud Infrastructure Modernisation Services\n"
            "RFP Reference: GOV-2026-ICT-0042\n"
            "Submission Deadline: 30 June 2026\n\n"
            "1. Technical Requirements\n"
            "REQ-T1: The vendor must demonstrate proven experience migrating Windows Server workloads "
            "to Azure IaaS and PaaS. Minimum 3 completed migrations of 500+ VMs with reference contacts.\n"
            "REQ-T2: The solution must achieve 99.95% uptime SLA for Tier-1 workloads with MTTR under 30 minutes.\n"
            "REQ-T3: All data must remain within Australian Azure regions. Vendor must hold ISO 27001.\n"
            "REQ-T4: Zero-downtime migration plan required using Azure Site Recovery. Cutover windows max 4 hours.\n\n"
            "2. Commercial Requirements\n"
            "REQ-C1: Fixed-price schedule for Years 1-3 with capped annual escalation of CPI + 2%.\n"
            "REQ-C2: Payment schedule tied to project milestones. Net-30 payment terms.\n\n"
            "3. Team and Governance\n"
            "REQ-G1: Named Project Manager with PMP/PRINCE2 and 7 years cloud migration experience required.\n"
            "REQ-G2: Monthly steering committee and fortnightly progress reports required.\n\n"
            "4. Legal\n"
            "REQ-L1: Compliance with Privacy Act 1988 and Notifiable Data Breaches scheme. "
            "24-hour breach notification required.\n\n"
            "5. Evaluation Criteria\n"
            "Technical capability 40%, Price 30%, Team experience 20%, Risk management 10%."
        ),
        "pages": 12,
        "sections": [
            {
                "heading": "Technical Requirements",
                "content": (
                    "REQ-T1: Proven experience migrating Windows Server workloads to Azure IaaS and PaaS. "
                    "Minimum 3 completed migrations of 500+ VMs with reference contacts. "
                    "REQ-T2: 99.95% uptime SLA for Tier-1 workloads, MTTR under 30 minutes. "
                    "REQ-T3: Data residency in Australian Azure regions, ISO 27001 required. "
                    "REQ-T4: Zero-downtime migration using Azure Site Recovery, max 4-hour cutover windows."
                ),
            },
            {
                "heading": "Commercial Requirements",
                "content": (
                    "REQ-C1: Fixed-price 3-year schedule, CPI+2% annual cap, all-inclusive pricing. "
                    "REQ-C2: Milestone-based payment schedule, net-30 terms."
                ),
            },
            {
                "heading": "Team and Governance",
                "content": (
                    "REQ-G1: Named PM with PMP/PRINCE2 and 7 years cloud migration experience. "
                    "REQ-G2: Monthly steering committee, fortnightly written progress reports."
                ),
            },
            {
                "heading": "Legal and Compliance",
                "content": (
                    "REQ-L1: Privacy Act 1988 compliance, Notifiable Data Breaches scheme, "
                    "24-hour breach notification."
                ),
            },
            {
                "heading": "Evaluation Criteria",
                "content": (
                    "Technical capability 40%, Price 30%, Team experience 20%, Risk management 10%. "
                    "Proposals below 60% overall will not proceed."
                ),
            },
        ],
    }
