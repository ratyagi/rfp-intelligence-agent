"""Bid Decision Report — the pipeline's reasoning trace, made first-class.

For every requirement, the report records the full decision chain:
evidence considered → score and confidence → citation verification outcome
→ action required. It is rendered three ways:

  1. streamed to the console while the pipeline runs (build_report logs it),
  2. appended to the proposal DOCX as an appendix (append_report_to_docx),
  3. written as a machine-readable JSON artifact (the Review stage saves it).

The bid recommendation is deterministic, derived from the priority-weighted
win probability — the model never decides whether to bid.
"""
import logging

from docx import Document
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

RECOMMENDATION_BANDS = [
    (70, "BID", "Strong evidence position across weighted requirements."),
    (50, "BID WITH CONDITIONS", "Viable position, but close the flagged gaps before submitting."),
    (0, "REVIEW BID DECISION", "Material gaps on weighted requirements — escalate before committing."),
]

_GREY = RGBColor(0x44, 0x44, 0x44)
_RED = RGBColor(0xCC, 0x00, 0x00)
_AMBER = RGBColor(0xB3, 0x6B, 0x00)
_GREEN = RGBColor(0x1E, 0x7A, 0x1E)
_SCORE_COLORS = {"COVERED": _GREEN, "PARTIAL": _AMBER, "GAP": _RED}


def build_report(verified_draft: dict, evidence_map: dict) -> dict:
    """Assemble the decision chain for every requirement and the bid summary."""
    requirements = verified_draft.get("requirements", [])
    verification = verified_draft.get("verification") or {}
    win_probability = verified_draft.get("win_probability", 0)

    recommendation, rationale = _recommend(win_probability)
    entries = [_entry(req, evidence_map.get(req["id"], [])) for req in requirements]

    report = {
        "rfp_title": verified_draft.get("rfp_title", ""),
        "win_probability": win_probability,
        "recommendation": recommendation,
        "recommendation_rationale": rationale,
        "counts": {
            "total": len(requirements),
            "covered": sum(1 for r in requirements if r.get("score") == "COVERED"),
            "partial": sum(1 for r in requirements if r.get("score") == "PARTIAL"),
            "gap": sum(1 for r in requirements if r.get("score") == "GAP"),
        },
        "citations": {
            "total": verification.get("citations_total", 0),
            "verified": verification.get("citations_verified", 0),
            "stripped": verification.get("citations_stripped", 0),
        },
        "requirements": entries,
    }

    _stream(report)
    return report


def _entry(req: dict, evidence: list) -> dict:
    verification = req.get("verification") or {}
    actions = []
    if req.get("gap_note"):
        actions.append(req["gap_note"])

    return {
        "id": req["id"],
        "requirement": (req.get("text") or "")[:200],
        "priority": req.get("priority", "medium"),
        "category": req.get("category", "other"),
        "evidence_considered": [
            {"doc_id": e["doc_id"], "title": e["title"], "retrieval_score": e.get("score")}
            for e in evidence
        ],
        "decision": {
            "score": req.get("score"),
            "scorer_confidence": req.get("confidence"),
            "citations_verified": verification.get("verified", []),
            "citations_stripped": verification.get("stripped", []),
        },
        "action_required": actions[0] if actions else None,
    }


def _recommend(win_probability: int) -> tuple[str, str]:
    for threshold, label, rationale in RECOMMENDATION_BANDS:
        if win_probability >= threshold:
            return label, rationale
    return RECOMMENDATION_BANDS[-1][1], RECOMMENDATION_BANDS[-1][2]


def _stream(report: dict) -> None:
    """Log the decision chain so the reasoning is visible during the run."""
    counts = report["counts"]
    citations = report["citations"]
    logger.info(
        f"BidReport: {report['recommendation']} — win probability "
        f"{report['win_probability']}% ({counts['covered']} covered, "
        f"{counts['partial']} partial, {counts['gap']} gap; "
        f"citations {citations['verified']}/{citations['total']} verified)"
    )
    for entry in report["requirements"]:
        evidence = ", ".join(e["doc_id"] for e in entry["evidence_considered"]) or "none"
        decision = entry["decision"]
        confidence = decision.get("scorer_confidence")
        confidence_str = f", confidence {confidence:.2f}" if confidence is not None else ""
        stripped = decision.get("citations_stripped") or []
        stripped_str = f", stripped {','.join(stripped)}" if stripped else ""
        logger.info(
            f"BidReport: {entry['id']} [{entry['priority']}/{entry['category']}] "
            f"evidence: {evidence} → {decision['score']}{confidence_str}{stripped_str}"
        )


def append_report_to_docx(docx_path: str, report: dict) -> None:
    """Append the Bid Decision Report as an appendix to the proposal DOCX."""
    doc = Document(docx_path)
    doc.add_page_break()

    # The proposal template has no built-in heading styles — style manually.
    title = doc.add_paragraph()
    title_run = title.add_run("Appendix — Bid Decision Report")
    title_run.bold = True
    title_run.font.size = Pt(16)

    counts = report["counts"]
    citations = report["citations"]
    summary = doc.add_paragraph()
    run = summary.add_run(
        f"Recommendation: {report['recommendation']} — "
        f"priority-weighted win probability {report['win_probability']}%. "
    )
    run.bold = True
    summary.add_run(
        f"{report['recommendation_rationale']} "
        f"Coverage: {counts['covered']} covered, {counts['partial']} partial, "
        f"{counts['gap']} gap of {counts['total']} requirements. "
        f"Citations verified: {citations['verified']}/{citations['total']}"
        + (f" ({citations['stripped']} stripped by the Verifier)." if citations["stripped"] else ".")
    )
    doc.add_paragraph(
        "This appendix is the pipeline's reasoning trace: for each requirement, "
        "the evidence retrieved from the knowledge base, the scoring decision "
        "with confidence, the citation verification outcome, and the action "
        "required from the bid team."
    ).runs[0].italic = True

    for entry in report["requirements"]:
        heading = doc.add_paragraph()
        head_run = heading.add_run(
            f"{entry['id']} ({entry['priority']} priority, {entry['category']}) — "
            f"{entry['decision']['score']}"
        )
        head_run.bold = True
        head_run.font.color.rgb = _SCORE_COLORS.get(entry["decision"]["score"], _GREY)

        _detail(doc, f"Requirement: {entry['requirement']}")
        evidence = entry["evidence_considered"]
        _detail(doc, "Evidence considered: " + (
            "; ".join(f"{e['doc_id']} ({e['title']})" for e in evidence)
            if evidence else "none — no relevant documents in the knowledge base"
        ))

        decision = entry["decision"]
        confidence = decision.get("scorer_confidence")
        decision_line = f"Decision: {decision['score']}"
        if confidence is not None:
            decision_line += f" (scorer confidence {confidence:.2f})"
        if decision.get("citations_verified"):
            decision_line += f"; citations verified: {', '.join(decision['citations_verified'])}"
        if decision.get("citations_stripped"):
            decision_line += f"; citations STRIPPED by Verifier: {', '.join(decision['citations_stripped'])}"
        _detail(doc, decision_line)

        if entry.get("action_required"):
            action = doc.add_paragraph()
            action_run = action.add_run(f"ACTION REQUIRED: {entry['action_required']}")
            action_run.font.color.rgb = _RED
            action_run.font.size = Pt(9)

    doc.save(docx_path)


def _detail(doc, text: str) -> None:
    para = doc.add_paragraph(text)
    para.runs[0].font.size = Pt(9)
    para.runs[0].font.color.rgb = _GREY
