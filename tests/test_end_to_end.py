"""End-to-end pipeline test — runs entirely in STUB_MODE=true."""
import os
import pytest

os.environ["STUB_MODE"] = "true"


@pytest.fixture(scope="module")
def pipeline_result():
    from agents.orchestrator import run_pipeline
    return run_pipeline({
        "rfp_source": "demo/sample_rfp.pdf",
        "company_name": "Contoso Cloud Solutions",
        "rfp_title": "GOV-2026-ICT-0042 Cloud Infrastructure Modernisation",
        "submission_deadline": "30 June 2026",
    })


def test_pipeline_status_complete(pipeline_result):
    assert pipeline_result["status"] == "complete", (
        f"Expected 'complete', got '{pipeline_result['status']}'. "
        f"Errors: {pipeline_result.get('errors', [])}"
    )


def test_pipeline_docx_path_exists(pipeline_result):
    docx_path = pipeline_result.get("docx_path")
    assert docx_path is not None, "docx_path is None"
    assert os.path.exists(docx_path), f"DOCX file not found at: {docx_path}"


def test_pipeline_docx_is_valid(pipeline_result):
    from docx import Document
    docx_path = pipeline_result["docx_path"]
    doc = Document(docx_path)
    assert len(doc.paragraphs) > 0, "DOCX has no paragraphs"


def test_pipeline_win_probability_in_range(pipeline_result):
    wp = pipeline_result.get("win_probability")
    assert wp is not None, "win_probability is None"
    assert isinstance(wp, int), f"win_probability must be int, got {type(wp)}"
    assert 0 <= wp <= 100, f"win_probability {wp} is out of range [0, 100]"


def test_pipeline_gap_count_nonnegative(pipeline_result):
    gap_count = pipeline_result.get("gap_count")
    assert gap_count is not None, "gap_count is None"
    assert gap_count >= 0, f"gap_count {gap_count} is negative"


def test_docx_contains_requirement_response(pipeline_result):
    from docx import Document
    doc = Document(pipeline_result["docx_path"])
    full_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + cell.text

    assert "REQ-" in full_text, "No requirement IDs found in document"


def test_docx_contains_confidential_footer(pipeline_result):
    from docx import Document
    doc = Document(pipeline_result["docx_path"])
    footer_text = " ".join(
        p.text for section in doc.sections for p in section.footer.paragraphs
    )
    assert "CONFIDENTIAL" in footer_text, (
        f"CONFIDENTIAL not found in footer. Footer text: '{footer_text}'"
    )
    assert "DRAFT" in footer_text, "DRAFT not found in footer"


def test_docx_no_unfilled_placeholders(pipeline_result):
    from docx import Document
    doc = Document(pipeline_result["docx_path"])
    full_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + cell.text

    assert "{{" not in full_text, (
        "Found unfilled {{placeholder}} tags in the document"
    )


def test_teams_card_posted(pipeline_result):
    assert pipeline_result.get("teams_card_posted") is True, (
        "Teams card was not posted (expected True in STUB mode)"
    )
